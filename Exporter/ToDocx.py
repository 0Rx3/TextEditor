from docx import Document
from docx.shared import Pt


def docx_run_qt(qt_doc, blockId, docx_paragraph):
    block = qt_doc.findBlockByNumber(blockId)
    font = block.charFormat().font()
    run = docx_paragraph.add_run()
    run.font.name = font.family()
    run.font.size = Pt(font.pointSize())
    run.bold = font.bold()
    run.italic = font.italic()
    run.underline = font.underline() != 0
    run.text = block.text()
    return run


def convert_qt_to_docx(qt_document, docx_filename):
    # create a new docx document
    docx_document = Document()

    for index in range(1, qt_document.blockCount()):
        p = docx_document.add_paragraph()
        docx_run_qt(qt_document, index, p)

    # Apply format again to the first paragraph
    block = qt_document.findBlockByNumber(0)
    font = block.charFormat().font()
    p = docx_document.paragraphs[0]
    run_again = p.add_run()
    run_again.font.name = font.family()
    run_again.font.size = Pt(font.pointSize())
    run_again.bold = font.bold()
    run_again.italic = font.italic()
    run_again.underline = font.underline() != 0

    docx_document.save(docx_filename)

    # Not finished because docx is an awful format.
